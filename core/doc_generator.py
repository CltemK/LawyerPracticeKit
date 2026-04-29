"""
Word文档生成器
将OCR识别的截图按时间排序后生成Word文档
"""
import math
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image as PILImage

from .time_parser import TimeParser


class DocGenerator:
    """Word文档生成器"""

    # A4页面尺寸（英寸）
    PAGE_WIDTH = 8.27
    PAGE_HEIGHT = 11.69
    MARGIN = 0.5

    def __init__(self, images_per_page=4, show_timestamp=True):
        self.images_per_page = max(1, min(9, images_per_page))
        self.show_timestamp = show_timestamp
        self.grid_rows, self.grid_cols = self._calc_grid(self.images_per_page)

    @staticmethod
    def _calc_grid(n):
        """根据每页图片数计算网格布局 (rows, cols)"""
        if n <= 0:
            return 1, 1
        cols = math.ceil(math.sqrt(n))
        rows = math.ceil(n / cols)
        return rows, cols

    def _calc_image_size(self, img_path, cell_w, cell_h):
        """计算保持纵横比的图片尺寸，返回 (width_inches, height_inches)"""
        with PILImage.open(img_path) as img:
            iw, ih = img.size
        ratio = iw / ih

        # 以单元格宽度为基准算高度
        h_by_w = cell_w / ratio
        if h_by_w <= cell_h:
            return cell_w, h_by_w
        # 超高，以单元格高度为基准算宽度
        return cell_h * ratio, cell_h

    def generate_from_ocr_results(self, ocr_results, output_path, time_parser=None, title=None):
        """
        从OCR结果列表生成Word文档
        ocr_results: [{'image_path': str, 'text': str, 'success': bool}, ...]
        """
        if time_parser is None:
            time_parser = TimeParser()

        # 为每张图片提取时间戳并排序
        items = []
        for result in ocr_results:
            image_path = result.get('image_path', '')
            text = result.get('text', '')
            timestamps = time_parser.extract_timestamps(text) if text else []
            time_range = None
            if timestamps:
                time_range = (timestamps[0][0], timestamps[-1][0])
            items.append({
                'image_path': image_path,
                'time_range': time_range,
                'earliest': timestamps[0][0] if timestamps else None,
            })

        # 有时间的按时间排，没时间的按文件名排
        items.sort(key=lambda x: (x['earliest'] is None, x['earliest'] or datetime.min, x['image_path']))

        return self._create_document(items, output_path, title or '微信聊天记录整理')

    def _create_document(self, items, output_path, title):
        """创建Word文档"""
        doc = Document()

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Inches(self.MARGIN)
            section.bottom_margin = Inches(self.MARGIN)
            section.left_margin = Inches(self.MARGIN)
            section.right_margin = Inches(self.MARGIN)

        # 可用区域（全部给图片，无标题占用）
        avail_w = self.PAGE_WIDTH - 2 * self.MARGIN
        avail_h = self.PAGE_HEIGHT - 2 * self.MARGIN
        # 预留表格边框/行间距/段落间距的损耗
        padding_loss = 0.3
        timestamp_h = 0.3 if self.show_timestamp else 0
        content_h = avail_h - padding_loss - timestamp_h * self.grid_rows

        cell_w = avail_w / self.grid_cols - 0.1
        cell_h = content_h / self.grid_rows - 0.1

        # 分页写入
        for page_idx in range(0, len(items), self.images_per_page):
            page_items = items[page_idx:page_idx + self.images_per_page]

            if page_idx > 0:
                doc.add_page_break()

            # 计算本页实际行列数
            n = len(page_items)
            rows = math.ceil(n / self.grid_cols)
            cols = self.grid_cols

            # 创建表格（图片行+时间戳行交替）
            table_rows = rows * 2 if self.show_timestamp else rows
            table = doc.add_table(rows=table_rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 设置表格单元格边距为0，最大化图片空间
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

            for idx, item in enumerate(page_items):
                r = idx // cols
                c = idx % cols

                img_cell = table.cell(r * 2 if self.show_timestamp else r, c)
                img_para = img_cell.paragraphs[0]
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                img_path = item['image_path']
                if img_path and os.path.exists(img_path):
                    w, h = self._calc_image_size(img_path, cell_w, cell_h)
                    run = img_para.add_run()
                    run.add_picture(img_path, width=Inches(w), height=Inches(h))

                # 时间戳
                if self.show_timestamp:
                    ts_cell = table.cell(r * 2 + 1, c)
                    ts_para = ts_cell.paragraphs[0]
                    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if item['time_range']:
                        ts_text = TimeParser.format_time_range(item['time_range'][0], item['time_range'][1])
                    else:
                        ts_text = '未识别时间'

                    run = ts_para.add_run(ts_text)
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.save(output_path)
        return output_path
